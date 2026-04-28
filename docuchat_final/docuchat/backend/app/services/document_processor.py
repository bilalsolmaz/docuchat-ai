import os
import re
from typing import List

import fitz  # PyMuPDF - PDF için
from docx import Document as DocxDocument


# ──────────────────────────────────────────────
# METİN ÇIKARMA
# ──────────────────────────────────────────────

def extract_text(file_path: str, file_name: str) -> str:
    """Dosya uzantısına göre metni çıkarır."""
    ext = os.path.splitext(file_name)[1].lower()

    if ext == ".txt":
        return _extract_txt(file_path)
    elif ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext in (".doc", ".docx"):
        return _extract_docx(file_path)
    else:
        raise ValueError(f"Desteklenmeyen dosya formatı: {ext}")


def _extract_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _extract_pdf(path: str) -> str:
    """PyMuPDF ile PDF'ten metin çıkarır (çok daha güvenilir)."""
    text_parts = []
    with fitz.open(path) as doc:
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(f"[Sayfa {page_num}]\n{page_text}")
    return "\n\n".join(text_parts)


def _extract_docx(path: str) -> str:
    doc = DocxDocument(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Tablolardan da metin al
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs)


# ──────────────────────────────────────────────
# METİN TEMİZLEME
# ──────────────────────────────────────────────

def clean_text(text: str) -> str:
    """Gereksiz boşluk ve karakterleri temizler."""
    text = re.sub(r"\n{3,}", "\n\n", text)          # Üçten fazla boş satırı ikiye indir
    text = re.sub(r" {2,}", " ", text)               # Çok boşlukları teke indir
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)  # Kontrol karakterleri
    return text.strip()


# ──────────────────────────────────────────────
# CHUNKING (PARÇALAMA)
# ──────────────────────────────────────────────

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """
    Metni örtüşen parçalara böler.
    - chunk_size: her parçanın karakter sayısı
    - overlap: iki parça arasındaki örtüşme miktarı (bağlam kaybını önler)
    """
    if not text.strip():
        return []

    chunks = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = min(start + chunk_size, text_len)

        # Cümle sınırında bitirmeye çalış (daha doğal parçalama)
        if end < text_len:
            for sep in (". ", ".\n", "\n\n", "\n", " "):
                bp = text.rfind(sep, start + chunk_size // 2, end)
                if bp != -1:
                    end = bp + len(sep)
                    break

        chunk = text[start:end].strip()
        if len(chunk) > 50 or (len(chunk) > 0 and len(chunks) == 0):  # Eğer hiç chunk yoksa küçüğe de izin ver
            chunks.append(chunk)

        if end >= text_len:
            break

        start = end - overlap

    return chunks
